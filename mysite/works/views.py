from __future__ import annotations

from django.contrib import messages
from django.contrib.auth import get_user_model
from django.db import transaction
from django.db.models import Max, Q
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView
from django.core.files.base import ContentFile
from django.conf import settings
from django.http import FileResponse, JsonResponse, HttpResponseForbidden
from django.views.decorators.csrf import csrf_exempt

from users.rbac import require_user_types
from profiles.course_utils import student_course_bucket
from profiles.models import Department, DirectionOfStudy, Student, Teacher
from profiles.seat_sync import sync_teacher_course_counts
from works.forms import (
    AssignTopicForm,
    DepartmentAnnouncementForm,
    PlanStageForm,
    TeacherWorkTopicForm,
    WorkDocxUploadForm,
)
from works.models import (
    IndividualPlan,
    Notification,
    PlanStage,
    StudentTopicProposal,
    Tag,
    Topic,
    TopicApplication,
    Work,
)
from works.student_work_type import type_of_work_label_for_student
from works.plan_utils import (
    flatten_plan_stages,
    leaf_stages_queryset,
    plan_progress_percent,
    reconcile_parent_done,
    seed_default_stages,
)

from io import BytesIO
from pathlib import Path
import hashlib
import json
import urllib.request
from urllib.parse import urlparse, urlunparse
import base64
import hmac
from types import SimpleNamespace


def _notify(recipient, title: str, body: str = "", url: str = "") -> None:
    Notification.objects.create(recipient=recipient, title=title, body=body, url=url)


def _work_single_version(work: Work):
    if not work.docx_file:
        return None
    return SimpleNamespace(
        file=work.docx_file,
        uploaded_at=work.docx_updated_at,
    )


def _user_display_name(user) -> str:
    return (user.get_full_name() or "").strip() or user.username


def _student_direction(student: Student) -> DirectionOfStudy | None:
    if not student.group or not student.group.direction_of_study_id:
        return None
    return student.group.direction_of_study


def _can_access_work(user, work: Work) -> bool:
    if not user.is_authenticated:
        return False
    ut = getattr(user, "user_type", None)
    if ut == "department_admin":
        dept_id = getattr(getattr(user, "teacher", None), "department_id", None)
        if dept_id and work.scientific_director_id:
            if work.scientific_director.department_id != dept_id:
                return False
        return True
    if hasattr(user, "student") and work.author_id == user.student.id:
        return True
    if hasattr(user, "teacher") and work.scientific_director_id == user.teacher.id:
        return True
    return False


def _student_primary_work(request):
    st = request.user.student
    w = st.works.order_by("-created_at").first()
    if not w:
        w = Work.objects.create(
            author=st,
            name=f"Работа ({st.user.get_full_name() or st.user.username})",
            scientific_director=st.scientific_director,
        )
    return w


def _ensure_work_document(work: Work) -> None:
    """
    Ensure minimal related infrastructure exists for a work.

    Currently this means ensuring an IndividualPlan record exists so the work
    can immediately be tracked by stages and progress.
    """

    IndividualPlan.objects.get_or_create(work=work)


def _generate_initial_docx(work: Work, uploaded_by) -> None:
    """
    Generate a minimal .docx file for the approved topic.
    If python-docx is not installed, silently skip (prototype-friendly).
    """
    if work.docx_file and (work.docx_file.name or "").lower().endswith(".docx"):
        return
    try:
        from docx import Document                
    except Exception:
        return

    d = Document()
    d.add_heading(work.name, level=1)
    if work.description:
        d.add_paragraph(work.description)
    d.add_paragraph("")
    d.add_paragraph("Текст работы:")

    bio = BytesIO()
    d.save(bio)
    bio.seek(0)

    if work.docx_file:
        try:
            work.docx_file.delete(save=False)
        except Exception:
            pass
    work.docx_updated_at = timezone.now()
    work.docx_file.save(f"work_{work.id}_initial.docx", ContentFile(bio.read()), save=True)


class SearchWorksView(LoginRequiredMixin, ListView):
    model = Work
    template_name = "works/home.html"
    extra_context = {"title": "Архив работ"}

    def get_queryset(self):
        qs = Work.objects.filter(status=Work.LifecycleStatus.DEFENDED).select_related(
            "author__user", "scientific_director__user"
        )
        q = self.request.GET.get("q")
        year = self.request.GET.get("year")
        teacher_id = self.request.GET.get("teacher_id")
        if q:
            qs = qs.filter(Q(name__icontains=q) | Q(description__icontains=q))
        if year:
            try:
                y = int(year)
                qs = qs.filter(created_at__year=y)
            except ValueError:
                pass
        if teacher_id:
            qs = qs.filter(scientific_director_id=teacher_id)
        for tid in self.request.GET.getlist("tag"):
            try:
                qs = qs.filter(tags__id=int(tid))
            except ValueError:
                pass
        return qs.distinct()

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx["teachers"] = Teacher.objects.select_related("user").order_by(
            "user__last_name", "user__first_name", "user__username"
        )
        ctx["years"] = Work.objects.filter(status=Work.LifecycleStatus.DEFENDED).dates(
            "created_at", "year", order="DESC"
        )
        ctx["all_tags"] = Tag.objects.order_by("name")[:500]
        works = ctx.get("object_list") or []
        with_docx: set[int] = set()
        for w in works:
            try:
                name = (w.docx_file.name or "").lower() if w.docx_file else ""
            except Exception:
                name = ""
            if name.endswith(".docx"):
                with_docx.add(w.id)
        ctx["archive_work_ids_with_docx"] = with_docx
        return ctx


class ShowTeachersView(LoginRequiredMixin, ListView):
    model = Teacher
    template_name = "works/list_teachers.html"

    def get_queryset(self):
        return Teacher.objects.select_related("user").order_by("department", "user__first_name")

    def get_context_data(self, **kwargs):
        from profiles.course_utils import CourseBucket
        from profiles.models import TeacherCourseSeat

        ctx = super().get_context_data(**kwargs)
        directions = list(DirectionOfStudy.objects.order_by("level_of_training", "direction_code"))
        ctx["directions"] = directions
        ctx["bachelor_directions"] = [d for d in directions if d.level_of_training == DirectionOfStudy.BACHELORS_DEGREE]
        ctx["master_directions"] = [d for d in directions if d.level_of_training == DirectionOfStudy.MASTERS_DEGREE]
        ctx["faculty"] = list(Department.objects.order_by("name"))

        ut = getattr(self.request.user, "user_type", None)
        if ut == "student":
                                                                                         
                                                                                       
            buckets = tuple(b for b in CourseBucket.ALL if b != CourseBucket.M2)
            bucket_choices = tuple((code, label) for (code, label) in CourseBucket.CHOICES if code != CourseBucket.M2)
        else:
            buckets = CourseBucket.ALL
            bucket_choices = CourseBucket.CHOICES

        ctx["course_buckets"] = bucket_choices
        seat_by_teacher: dict[int, dict[str, tuple[int, int]]] = {}
        for s in TeacherCourseSeat.objects.all():
            seat_by_teacher.setdefault(s.teacher_id, {})[s.bucket] = (s.current_amount, s.max_amount)
        teacher_cells = {}
        for t in ctx["object_list"]:
            teacher_cells[t.id] = [seat_by_teacher.get(t.id, {}).get(b, (0, 5)) for b in buckets]
        ctx["teacher_cells"] = teacher_cells
        return ctx


class ShowStudentsView(LoginRequiredMixin, ListView):
    model = Student
    template_name = "works/add_student.html"

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        if getattr(self.request.user, "user_type", None) == "teacher":
            teacher = self.request.user.teacher
            apps = (
                TopicApplication.objects.filter(status=TopicApplication.Status.PENDING, topic__creator=teacher)
                .select_related("student__user", "topic", "student__group", "student__group__direction_of_study")
                .order_by("-created_at")
            )
            props = (
                StudentTopicProposal.objects.filter(teacher=teacher, status=StudentTopicProposal.Status.PENDING)
                .select_related("student__user", "student__group", "student__group__direction_of_study")
                .order_by("-created_at")
            )
            ctx["applications"] = apps
            ctx["topic_proposals"] = props
        teacher_qs = Teacher.objects.all()
        if getattr(self.request.user, "user_type", None) == "department_admin":
            dept_id = getattr(getattr(self.request.user, "teacher", None), "department_id", None)
            if dept_id:
                teacher_qs = teacher_qs.filter(department_id=dept_id)
        ctx["teachers"] = teacher_qs
        return ctx

    def get_queryset(self):
        query = self.request.GET.get("q")
        ut = getattr(self.request.user, "user_type", None)
        if ut == "teacher":
            teacher = self.request.user.teacher
            app_ids = TopicApplication.objects.filter(
                status=TopicApplication.Status.PENDING, topic__creator=teacher
            ).values_list("student_id", flat=True)
            prop_ids = StudentTopicProposal.objects.filter(
                teacher=teacher, status=StudentTopicProposal.Status.PENDING
            ).values_list("student_id", flat=True)
            student_ids = set(app_ids) | set(prop_ids)
            qs = Student.objects.filter(id__in=student_ids).select_related("user", "group", "group__direction_of_study")
        else:
            qs = Student.objects.filter(scientific_director__isnull=True).select_related(
                "user", "group", "group__direction_of_study"
            )
        if query is not None:
            return qs.filter(Q(user__first_name__icontains=query))
        return qs


@require_user_types("teacher", "department_admin")
def check_works(request):
    user_type = request.user.user_type

    if request.method == "POST":
        student = Student.objects.get(id=request.POST.get("student_id"))
        if user_type == "teacher":
            teacher = Teacher.objects.get(user=request.user)
        else:
            teacher_id = request.POST.get("teacher_id")
            teacher = get_object_or_404(Teacher, id=teacher_id)
            dept_id = getattr(getattr(request.user, "teacher", None), "department_id", None)
            if dept_id and teacher.department_id != dept_id:
                return redirect("add_student")

        if user_type == "teacher":
            proposal_id = request.POST.get("proposal_id")
            if proposal_id:
                prop = get_object_or_404(
                    StudentTopicProposal,
                    id=proposal_id,
                    teacher=teacher,
                    student=student,
                    status=StudentTopicProposal.Status.PENDING,
                )
                if not _has_available_seat(teacher, student):
                    messages.error(
                        request,
                        "Нет свободных мест по курсу студента для этого преподавателя.",
                    )
                    return redirect("add_student")
                prop.status = StudentTopicProposal.Status.APPROVED
                prop.decided_at = timezone.now()
                prop.save(update_fields=["status", "decided_at"])
                student.scientific_director = teacher
                student.save()
                sync_teacher_course_counts(teacher)
                buck = student_course_bucket(student)
                if buck:
                    _reject_pending_when_no_seats(teacher, buck)
                tw = type_of_work_label_for_student(student)
                work, _created = Work.objects.get_or_create(
                    author=student,
                    defaults={
                        "scientific_director": teacher,
                        "status": Work.LifecycleStatus.NOT_READY,
                        "name": prop.title,
                        "description": prop.description,
                        "type_of_work": tw,
                    },
                )
                work.name = prop.title
                work.description = prop.description
                work.scientific_director = teacher
                work.type_of_work = tw
                work.save()
                _reject_other_pending_for_student(student, teacher.id)
                _notify(student.user, "Решение по предложенной теме", f"Утверждена: {prop.title}", "")
                _ensure_work_document(work)
                _generate_initial_docx(work, uploaded_by=request.user)
            else:
                app = (
                    TopicApplication.objects.filter(
                        status=TopicApplication.Status.PENDING, student=student, topic__creator=teacher
                    )
                    .select_related("topic")
                    .order_by("-created_at")
                    .first()
                )
                if not app:
                    messages.error(request, "Нет ожидающей заявки на тему для этого студента.")
                    return redirect("add_student")
                if not _has_available_seat(teacher, student):
                    messages.error(
                        request,
                        "Нет свободных мест по курсу студента для этого преподавателя.",
                    )
                    return redirect("add_student")
                app.status = TopicApplication.Status.APPROVED
                app.decided_at = timezone.now()
                app.save(update_fields=["status", "decided_at"])
                student.scientific_director = teacher
                student.save()
                sync_teacher_course_counts(teacher)
                buck = student_course_bucket(student)
                if buck:
                    _reject_pending_when_no_seats(teacher, buck)
                tw = app.topic.get_work_kind_display()
                work, _created = Work.objects.get_or_create(
                    author=student,
                    defaults={
                        "scientific_director": teacher,
                        "status": Work.LifecycleStatus.NOT_READY,
                        "name": app.topic.title,
                        "description": app.topic.description,
                        "type_of_work": tw,
                    },
                )
                work.name = app.topic.title
                work.description = app.topic.description
                work.scientific_director = teacher
                work.type_of_work = tw
                work.save()
                _reject_other_pending_applications_for_topic(app.topic, student.id)
                _reject_other_pending_for_student(student, teacher.id)
                _notify(
                    student.user,
                    "Решение по заявке на тему",
                    f"{app.get_status_display()}: {app.topic.title}",
                    "",
                )
                _ensure_work_document(work)
                _generate_initial_docx(work, uploaded_by=request.user)
        else:
            if not _has_available_seat(teacher, student):
                messages.error(request, "Нет свободных мест по курсу студента для этого преподавателя.")
                return redirect("add_student")
            student.scientific_director = teacher
            student.save()
            sync_teacher_course_counts(teacher)
            buck = student_course_bucket(student)
            if buck:
                _reject_pending_when_no_seats(teacher, buck)
            tw = type_of_work_label_for_student(student)
            work, _created = Work.objects.get_or_create(
                author=student,
                defaults={
                    "scientific_director": teacher,
                    "status": Work.LifecycleStatus.NOT_READY,
                    "name": f"Работа: {student.user.get_full_name() or student.user.username}",
                    "type_of_work": tw,
                },
            )
            if work.scientific_director_id != teacher.id:
                work.scientific_director = teacher
            if not (work.type_of_work or "").strip():
                work.type_of_work = tw
            work.save()

    if user_type == "teacher":
        teacher = Teacher.objects.get(user=request.user)
        students = Student.objects.filter(scientific_director=teacher)
        works = Work.objects.filter(author__in=students).select_related("author", "scientific_director")
    else:
        teacher_qs = Teacher.objects.all()
        dept = getattr(getattr(request.user, "teacher", None), "department", "") or ""
        if dept:
            teacher_qs = teacher_qs.filter(department=dept)
        students = Student.objects.filter(scientific_director__in=teacher_qs)
        works = Work.objects.filter(author__in=students).select_related("author", "scientific_director")

    work_by_student = {w.author_id: w for w in works if w.author_id}
    return render(
        request,
        "works/check_works.html",
        {"students": students, "works": works, "work_by_student": work_by_student},
    )


@require_user_types("student")
def write_work(request):
    work = _student_primary_work(request)
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    student = request.user.student
    plan = IndividualPlan.objects.filter(work=work).first()
    if plan:
        seed_default_stages(plan)
    plan_flat = flatten_plan_stages(plan) if plan else []
    progress, _, _ = plan_progress_percent(plan) if plan else (0, 0, 0)
    if request.method == "POST" and "save_tags" in request.POST:
        tag_objs = []
        for tid in request.POST.getlist("tag"):
            try:
                tag_objs.append(Tag.objects.get(id=int(tid)))
            except (ValueError, Tag.DoesNotExist):
                pass
        raw = request.POST.get("tags_csv", "")
        for name in [t.strip() for t in raw.replace(";", ",").split(",") if t.strip()][:200]:
            t, _ = Tag.objects.get_or_create(name=name[:64])
            tag_objs.append(t)
        work.tags.set(tag_objs)
        messages.success(request, "Теги сохранены.")
        return redirect("write_work")
    all_tags = Tag.objects.order_by("name")[:500]
    selected_ids = list(work.tags.values_list("id", flat=True))
    return render(
        request,
        "works/write_work.html",
        {
            "student": student,
            "primary_work": work,
            "plan": plan,
            "plan_flat": plan_flat,
            "progress": progress,
            "all_tags": all_tags,
            "selected_tag_ids": selected_ids,
        },
    )


@require_user_types("teacher", "department_admin")
def edit_work_topic(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if work.is_archived:
        return render(request, "works/forbidden.html", status=403)
    if getattr(request.user, "user_type", None) == "teacher":
        if not hasattr(request.user, "teacher") or work.scientific_director_id != request.user.teacher.id:
            return render(request, "works/forbidden.html", status=403)
    elif getattr(request.user, "user_type", None) == "department_admin":
        dept = getattr(getattr(request.user, "teacher", None), "department", "") or ""
        if dept and work.scientific_director and work.scientific_director.department != dept:
            return render(request, "works/forbidden.html", status=403)
    else:
        return render(request, "works/forbidden.html", status=403)

    if request.method == "POST":
        form = TeacherWorkTopicForm(request.POST, instance=work)
        if form.is_valid():
            form.save()
            if work.author:
                _notify(work.author.user, "Тема работы обновлена", work.name, f"/works/review_work/{work.id}/")
            return redirect("teacher_review_work", work_id=work.id)
    else:
        form = TeacherWorkTopicForm(instance=work)
    return render(request, "works/edit_work_topic.html", {"form": form, "work": work})


@require_user_types("department_admin")
def department_stats(request):
    teacher_qs = Teacher.objects.select_related("user").all()
    dept_id = getattr(getattr(request.user, "teacher", None), "department_id", None)
    if dept_id:
        teacher_qs = teacher_qs.filter(department_id=dept_id)
    teacher_qs = teacher_qs.order_by("user__first_name")

    teacher_sections = []
    for teacher in teacher_qs:
        students = Student.objects.filter(scientific_director=teacher).select_related("user", "group__direction_of_study")
        works = Work.objects.filter(author__in=students).select_related("author")
        work_by_student = {w.author_id: w for w in works if w.author_id}
        rows = []
        for st in students:
            w = work_by_student.get(st.id)
            plan = IndividualPlan.objects.filter(work=w).first() if w else None
            pct, _, _ = plan_progress_percent(plan) if plan else (0, 0, 0)
            rows.append({"student": st, "work": w, "progress": pct})
        teacher_sections.append({"teacher": teacher, "rows": rows})

    return render(
        request,
        "works/stats.html",
        {
            "teacher_sections": teacher_sections,
        },
    )


@require_user_types("department_admin")
def department_teachers_list(request):
    from profiles.course_utils import CourseBucket
    from profiles.models import TeacherCourseSeat
    from profiles.seat_sync import ensure_teacher_seats

    dept_id = getattr(getattr(request.user, "teacher", None), "department_id", None)
    if not dept_id:
        messages.error(request, "Не задана кафедра в профиле.")
        return redirect("home")
    teacher_qs = Teacher.objects.select_related("user").filter(department_id=dept_id).order_by("user__first_name")

    if request.method == "POST" and request.POST.get("update_limits"):
        for t in teacher_qs:
            ensure_teacher_seats(t)
            for b in CourseBucket.ALL:
                key = f"max_{t.id}_{b}"
                raw = request.POST.get(key)
                if raw is None:
                    continue
                try:
                    new_max = int(raw)
                except ValueError:
                    continue
                if new_max < 0:
                    continue
                rec = TeacherCourseSeat.objects.filter(teacher=t, bucket=b).first()
                if not rec:
                    continue
                if new_max < rec.current_amount:
                    messages.error(
                        request,
                        f"Лимит для {t.user.first_name} ({rec.get_bucket_display()}) не может быть ниже числа закреплённых студентов ({rec.current_amount}).",
                    )
                    return redirect("department_teachers")
                rec.max_amount = new_max
                rec.save(update_fields=["max_amount"])
        messages.success(request, "Лимиты сохранены.")
        return redirect("department_teachers")

    directions = list(DirectionOfStudy.objects.order_by("level_of_training", "direction_code"))
    seat_by_teacher: dict[int, dict[str, tuple[int, int]]] = {}
    for s in TeacherCourseSeat.objects.all():
        seat_by_teacher.setdefault(s.teacher_id, {})[s.bucket] = (s.current_amount, s.max_amount)
    teacher_cells = {}
    for t in teacher_qs:
        teacher_cells[t.id] = {b: seat_by_teacher.get(t.id, {}).get(b, (0, 5)) for b in CourseBucket.ALL}
                                                                                    
    course_buckets_table = tuple((code, label) for code, label in CourseBucket.CHOICES if code != CourseBucket.M2)
    return render(
        request,
        "works/department_teachers.html",
        {
            "teachers": teacher_qs,
            "directions": directions,
            "course_buckets": course_buckets_table,
            "bachelor_directions": [d for d in directions if d.level_of_training == DirectionOfStudy.BACHELORS_DEGREE],
            "master_directions": [d for d in directions if d.level_of_training == DirectionOfStudy.MASTERS_DEGREE],
            "teacher_cells": teacher_cells,
        },
    )


@require_user_types("department_admin")
def department_teacher_detail(request, teacher_id: int):
    teacher = get_object_or_404(Teacher, id=teacher_id)
    dept_id = getattr(getattr(request.user, "teacher", None), "department_id", None)
    if dept_id and teacher.department_id != dept_id:
        return redirect("department_teachers")
    students = Student.objects.filter(scientific_director=teacher).select_related("user", "group__direction_of_study")
    works = Work.objects.filter(author__in=students).select_related("author")
    work_by_student = {w.author_id: w for w in works if w.author_id}
    rows = []
    for st in students:
        w = work_by_student.get(st.id)
        plan = IndividualPlan.objects.filter(work=w).first() if w else None
        stages = list(plan.stages.order_by("order", "id")) if plan else []
        pct, _, _ = plan_progress_percent(plan) if plan else (0, 0, 0)
        rows.append({"student": st, "work": w, "stages": stages, "progress_pct": pct})
    return render(
        request,
        "works/department_teacher_detail.html",
        {"teacher": teacher, "rows": rows},
    )


@require_user_types("department_admin")
def department_send_announcement(request):
    """Массовое уведомление студентам выбранного курса кафедры и их научным руководителям (без дублей)."""
    dept_id = getattr(getattr(request.user, "teacher", None), "department_id", None)
    if not dept_id:
        messages.error(request, "Не задана кафедра в профиле.")
        return redirect("home")

    if request.method == "POST":
        form = DepartmentAnnouncementForm(request.POST)
        if form.is_valid():
            course = form.cleaned_data["course"]
            body = (form.cleaned_data.get("body") or "").strip()
            title = "Сообщение кафедры"
            url = reverse("notifications")

            students = Student.objects.filter(
                scientific_director__department_id=dept_id,
                group__course=course,
            ).select_related("user", "scientific_director__user")

            student_user_ids: set[int] = set()
            teacher_user_ids: set[int] = set()
            for st in students:
                if st.user_id:
                    student_user_ids.add(st.user_id)
                if st.scientific_director_id and st.scientific_director.user_id:
                    teacher_user_ids.add(st.scientific_director.user_id)

            User = get_user_model()
            admin_id = request.user.id
            notified: set[int] = set()
            n_st = 0
            for uid in student_user_ids:
                if uid == admin_id:
                    continue
                u = User.objects.filter(pk=uid).first()
                if u:
                    _notify(u, title, body, url)
                    notified.add(uid)
                    n_st += 1
            n_t = 0
            for uid in teacher_user_ids:
                if uid == admin_id or uid in notified:
                    continue
                u = User.objects.filter(pk=uid).first()
                if u:
                    _notify(u, title, body, url)
                    notified.add(uid)
                    n_t += 1
            messages.success(request, f"Создано уведомлений: студентам — {n_st}, преподавателям — {n_t}.")
            return redirect("department_send_announcement")
    else:
        form = DepartmentAnnouncementForm()

    return render(
        request,
        "works/department_send_announcement.html",
        {"form": form},
    )


@require_user_types("department_admin")
def assign_topic(request, student_id: int):
    student = get_object_or_404(Student, id=student_id)
    work = Work.objects.filter(author=student).first()
    if not work:
        work = Work(author=student, status=Work.LifecycleStatus.NOT_READY, name="Тема не задана")

    if request.method == "POST":
        form = AssignTopicForm(request.POST, instance=work, request_user=request.user)
        if form.is_valid():
            form.save()
            return redirect("check_works")
    else:
        form = AssignTopicForm(instance=work, request_user=request.user)

    return render(request, "works/assign_topic.html", {"student": student, "form": form})


@login_required
def work_detail(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if work.is_archived:
        v = _work_single_version(work)
        versions = [v] if v else []
        has_docx = bool(work.docx_file)
        plan = IndividualPlan.objects.filter(work=work).first()
        progress, _done, _total = plan_progress_percent(plan) if plan else (0, 0, 0)
        return render(
            request,
            "works/work_archive_view.html",
            {
                "work": work,
                "versions": versions,
                "has_docx": has_docx,
                "progress": progress,
            },
        )
    if not _can_access_work(request.user, work):
        return render(request, "works/forbidden.html", status=403)
    if hasattr(request.user, "student") and work.author_id == request.user.student.id:
        return redirect("write_work")
    if getattr(request.user, "user_type", None) in ("teacher", "department_admin"):
        return redirect("teacher_review_work", work_id=work.id)
    return redirect("home")


@require_user_types("teacher", "department_admin")
def teacher_review_work(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not _can_access_work(request.user, work):
        return render(request, "works/forbidden.html", status=403)
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    student = work.author
    if not student:
        return render(request, "works/forbidden.html", status=403)
    plan = IndividualPlan.objects.filter(work=work).first()
    if not plan and getattr(request.user, "user_type", None) in ("teacher", "department_admin"):
        plan = IndividualPlan.objects.create(work=work)
    if plan:
        seed_default_stages(plan)
    plan_flat = flatten_plan_stages(plan) if plan else []
    progress, _done, _total = plan_progress_percent(plan) if plan else (0, 0, 0)
    leaves = list(leaf_stages_queryset(plan)) if plan else []
    all_leaves_done = bool(leaves) and all(s.is_done for s in leaves)
    ut = getattr(request.user, "user_type", None)
    is_director = hasattr(request.user, "teacher") and work.scientific_director_id == request.user.teacher.id
    can_mark_ready = (
        ut == "teacher"
        and is_director
        and all_leaves_done
        and work.status == Work.LifecycleStatus.NOT_READY
    )
    can_mark_norm = ut == "department_admin" and work.status == Work.LifecycleStatus.READY
    norm_button_visible_admin = ut == "department_admin"
    can_mark_defended = ut == "teacher" and is_director and work.status == Work.LifecycleStatus.NORM_OK
    stage_form = PlanStageForm(plan=plan) if plan else PlanStageForm()
    return render(
        request,
        "works/teacher_review_work.html",
        {
            "work": work,
            "student": student,
            "plan": plan,
            "plan_flat": plan_flat,
            "stage_form": stage_form,
            "progress": progress,
            "can_mark_ready": can_mark_ready,
            "can_mark_norm": can_mark_norm,
            "norm_button_visible_admin": norm_button_visible_admin,
            "can_mark_defended": can_mark_defended,
        },
    )


@require_POST
@require_user_types("teacher")
def mark_work_ready(request, work_id: int):
    """Научный руководитель: все этапы выполнены → статус «готова» (ожидание нормоконтроля)."""
    work = get_object_or_404(Work, id=work_id)
    if not hasattr(request.user, "teacher") or work.scientific_director_id != request.user.teacher.id:
        return redirect("home")
    if not _can_access_work(request.user, work):
        return redirect("home")
    plan = IndividualPlan.objects.filter(work=work).first()
    leaves = list(leaf_stages_queryset(plan)) if plan else []
    if not leaves or not all(s.is_done for s in leaves):
        messages.error(request, "Все пункты плана (листья дерева) должны быть отмечены выполненными.")
        return redirect("teacher_review_work", work_id=work.id)
    if work.status != Work.LifecycleStatus.NOT_READY:
        messages.error(request, "Недопустимый переход статуса.")
        return redirect("teacher_review_work", work_id=work.id)
    work.status = Work.LifecycleStatus.READY
    work.save(update_fields=["status"])
    if work.author:
        _notify(
            work.author.user,
            "Работа отмечена как готовая",
            f"«{work.name}»: ожидается нормоконтроль на кафедре.",
            reverse("teacher_review_work", kwargs={"work_id": work.id}),
        )
    messages.success(request, "Работа отмечена как готовая к нормоконтролю.")
    return redirect("teacher_review_work", work_id=work.id)


@require_POST
@require_user_types("department_admin")
def mark_work_norm_ok(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.status != Work.LifecycleStatus.READY:
        messages.error(request, "Нормоконтроль можно отметить только для работы в статусе «готова».")
        return redirect("teacher_review_work", work_id=work.id)
    work.status = Work.LifecycleStatus.NORM_OK
    work.save(update_fields=["status"])
    if work.author:
        _notify(
            work.author.user,
            "Нормоконтроль пройден",
            f"«{work.name}»: можно готовиться к защите.",
            reverse("write_work"),
        )
    if work.scientific_director:
        st_name = "Студент"
        if work.author and work.author.user_id:
            st_name = work.author.user.get_full_name() or work.author.user.username
        _notify(
            work.scientific_director.user,
            "Нормоконтроль пройден",
            f"«{work.name}» ({st_name}).",
            reverse("teacher_review_work", kwargs={"work_id": work.id}),
        )
    messages.success(request, "Отмечено: нормоконтроль пройден.")
    return redirect("teacher_review_work", work_id=work.id)


@require_POST
@require_user_types("teacher")
def mark_work_defended(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not hasattr(request.user, "teacher") or work.scientific_director_id != request.user.teacher.id:
        return redirect("home")
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.status != Work.LifecycleStatus.NORM_OK:
        messages.error(request, "Защиту можно отметить только после нормоконтроля.")
        return redirect("teacher_review_work", work_id=work.id)
    _copy_final_docx_to_completed(work)
    work.status = Work.LifecycleStatus.DEFENDED
    work.completed_at = timezone.now()
    work.save(update_fields=["status", "completed_at"])
    if work.author:
        _notify(work.author.user, "Работа принята в архив", work.name, "")
    messages.success(request, "Работа защищена и перенесена в архив.")
    return redirect("work_detail", work_id=work.id)


@require_POST
@require_user_types("teacher", "department_admin")
def save_plan_stages(request, work_id: int):
    """Сохранение всех этапов плана одной отправкой формы (страница проверки работы)."""
    work = get_object_or_404(Work, id=work_id)
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.is_archived:
        messages.error(request, "Работа в архиве, план недоступен для изменений.")
        return redirect("work_detail", work_id=work.id)
    plan = IndividualPlan.objects.filter(work=work).first()
    if not plan:
        messages.error(request, "План не найден.")
        return redirect("teacher_review_work", work_id=work.id)

    stages = list(plan.stages.all())
    forms_to_save = []
    for stage in stages:
        prefix = f"stage_{stage.id}_"
        post = request.POST
        data = {
            "title": post.get(f"{prefix}title", "").strip(),
            "due_date": post.get(f"{prefix}due_date", "").strip() or None,
            "parent": post.get(f"{prefix}parent", "").strip() or None,
            "order": post.get(f"{prefix}order", "0"),
        }
        form = PlanStageForm(data, instance=stage, plan=plan)
        if not form.is_valid():
            messages.error(request, "Не удалось сохранить изменения: проверьте поля этапов плана.")
            return redirect("teacher_review_work", work_id=work.id)
        forms_to_save.append(form)

    with transaction.atomic():
        for form in forms_to_save:
            form.save()

    messages.success(request, "Изменения в плане сохранены.")
    if work.author:
        _notify(work.author.user, "План работы обновлён", work.name, "/works/write_work/")
    return redirect("teacher_review_work", work_id=work.id)


@require_POST
@require_user_types("teacher", "department_admin")
def edit_plan_stage(request, stage_id: int):
    stage = get_object_or_404(PlanStage, id=stage_id)
    work = stage.plan.work
    plan = stage.plan
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    form = PlanStageForm(request.POST, instance=stage, plan=plan)
    if form.is_valid():
        form.save()
        if work.author:
            _notify(work.author.user, "План работы обновлён", stage.title, "")
    return redirect("teacher_review_work", work_id=work.id)


@require_POST
@require_user_types("teacher", "department_admin")
def delete_plan_stage(request, stage_id: int):
    stage = get_object_or_404(PlanStage, id=stage_id)
    work = stage.plan.work
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    wid = work.id
    stage.delete()
    return redirect("teacher_review_work", work_id=wid)


@require_user_types("student")
def upload_version(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if work.author_id != request.user.student.id:
        return redirect("home")
    if work.is_archived:
        messages.error(request, "Работа в архиве, загрузка версий недоступна.")
        return redirect("work_detail", work_id=work.id)
    if request.method == "POST":
        form = WorkDocxUploadForm(request.POST, request.FILES, instance=work)
        if form.is_valid():
                                                                 
            if work.docx_file and "docx_file" in request.FILES:
                try:
                    work.docx_file.delete(save=False)
                except Exception:
                    pass
            saved = form.save(commit=False)
            saved.docx_updated_at = timezone.now()
            saved.save(update_fields=["docx_file", "docx_updated_at"])
            if work.scientific_director:
                _notify(
                    work.scientific_director.user,
                    "Новая версия работы",
                    f"{work.name}",
                    f"/works/work/{work.id}/",
                )
            return redirect("work_detail", work_id=work.id)
    else:
        form = WorkDocxUploadForm(instance=work)
    return render(request, "works/upload_version.html", {"form": form, "work": work})


@require_user_types("teacher", "department_admin")
def add_review(request, version_id: int):
                                                                                                  
    return redirect("home")


@login_required
def plan_detail(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not _can_access_work(request.user, work):
        return render(request, "works/forbidden.html", status=403)
    if work.is_archived and hasattr(request.user, "student") and work.author_id == request.user.student.id:
        return redirect("work_detail", work_id=work.id)

    plan = IndividualPlan.objects.filter(work=work).first()
    ut = getattr(request.user, "user_type", None)
    if not plan and (ut in {"teacher", "department_admin"}):
        plan = IndividualPlan.objects.create(work=work)
    if plan:
        seed_default_stages(plan)
    plan_flat = flatten_plan_stages(plan) if plan else []
    progress, _d, _t = plan_progress_percent(plan) if plan else (0, 0, 0)
    stage_form = PlanStageForm(plan=plan) if plan else PlanStageForm()
    return render(
        request,
        "works/plan_detail.html",
        {
            "work": work,
            "plan": plan,
            "plan_flat": plan_flat,
            "progress": progress,
            "stage_form": stage_form,
        },
    )


@require_user_types("teacher", "department_admin")
def add_plan_stage(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    plan, _ = IndividualPlan.objects.get_or_create(work=work)
    if request.method == "POST":
        form = PlanStageForm(request.POST, plan=plan)
        if form.is_valid():
            s = form.save(commit=False)
            s.plan = plan
            mx = plan.stages.aggregate(m=Max("order"))["m"]
            s.order = (mx or 0) + 1
            s.save()
            if work.author:
                _notify(work.author.user, "Новый этап плана", s.title, f"/works/review_work/{work.id}/")
    nxt = request.POST.get("next")
    if nxt:
        return redirect(nxt)
    return redirect("teacher_review_work", work_id=work.id)


@login_required
def mark_stage_ready(request, stage_id: int):
    stage = get_object_or_404(PlanStage, id=stage_id)
    work = stage.plan.work
    if not hasattr(request.user, "student") or work.author_id != request.user.student.id:
        return redirect("home")
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)
    stage.student_ready = True
    stage.ready_at = timezone.now()
    stage.save()
    if work.scientific_director:
        _notify(
            work.scientific_director.user,
            "Этап готов к проверке",
            stage.title,
            f"/works/work/{work.id}/plan/",
        )
    return redirect("plan_detail", work_id=work.id)


@require_user_types("teacher", "department_admin")
def approve_stage(request, stage_id: int):
    stage = get_object_or_404(PlanStage, id=stage_id)
    work = stage.plan.work
    if not _can_access_work(request.user, work):
        return redirect("home")
    if work.is_archived:
        return redirect("work_detail", work_id=work.id)

    if getattr(request.user, "user_type", None) == "teacher":
        if not hasattr(request.user, "teacher") or work.scientific_director_id != request.user.teacher.id:
            return redirect("home")

    stage.is_done = True
    stage.done_at = timezone.now()
    stage.save(update_fields=["is_done", "done_at"])
    reconcile_parent_done(stage)

    if work.author:
        _notify(work.author.user, "Этап плана подтверждён", stage.title, f"/works/review_work/{work.id}/")
    nxt = request.POST.get("next")
    if nxt:
        return redirect(nxt)
    return redirect("teacher_review_work", work_id=work.id)


@login_required
def notifications(request):
    if request.GET.get("mark_all"):
        Notification.objects.filter(recipient=request.user, is_read=False).update(is_read=True)
        return redirect("notifications")
    items = list(Notification.objects.filter(recipient=request.user)[:50])
    return render(request, "works/notifications.html", {"items": items})


def _copy_final_docx_to_completed(work: Work) -> bool:
    if not work.docx_file:
        return False
    dest_dir = Path(settings.MEDIA_ROOT) / "completed_works" / str(work.id)
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / "final.docx"
    with work.docx_file.open("rb") as src:
        dest.write_bytes(src.read())
    return True


def _b64url(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode("utf-8").rstrip("=")


def _jwt_hs256(payload: dict, secret: str) -> str:
    header = {"alg": "HS256", "typ": "JWT"}
    signing_input = f"{_b64url(json.dumps(header, separators=(',', ':')).encode('utf-8'))}.{_b64url(json.dumps(payload, separators=(',', ':')).encode('utf-8'))}"
    sig = hmac.new(secret.encode("utf-8"), signing_input.encode("utf-8"), hashlib.sha256).digest()
    return f"{signing_input}.{_b64url(sig)}"


def _has_docx(work: Work) -> bool:
    try:
        name = (work.docx_file.name or "").lower() if work.docx_file else ""
    except Exception:
        name = ""
    return bool(name.endswith(".docx"))


def _ds_reachable_url(abs_url: str) -> str:
    """
    Адреса localhost/127.0.0.1 в docker‑сети на стороне Document Server часто не открываются.
    Для URL, по которым OnlyOffice (в контейнере) скачивает DOCX, подменяем host на host.docker.internal.
    """
    p = urlparse(abs_url)
    host = (p.hostname or "").lower()
    if host in {"127.0.0.1", "localhost"}:
        port = p.port or 8000
        netloc = f"host.docker.internal:{port}"
        return urlunparse((p.scheme, netloc, p.path, p.params, p.query, p.fragment))
    return abs_url


def _onlyoffice_file_and_callback_urls(request, work: Work) -> tuple[str, str]:
    token = settings.ONLYOFFICE_CALLBACK_TOKEN
    if settings.ONLYOFFICE_APP_PUBLIC_URL:
        base = settings.ONLYOFFICE_APP_PUBLIC_URL.rstrip("/")
        file_url = f"{base}/works/work/{work.id}/docx/file/?token={token}"
        callback_url = f"{base}/works/onlyoffice/callback/{work.id}/?token={token}"
        return file_url, callback_url
    file_url = _ds_reachable_url(
        request.build_absolute_uri(f"/works/work/{work.id}/docx/file/") + f"?token={token}"
    )
    callback_url = _ds_reachable_url(
        request.build_absolute_uri(f"/works/onlyoffice/callback/{work.id}/") + f"?token={token}"
    )
    return file_url, callback_url


def onlyoffice_docx_file(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    token = request.GET.get("token")
    ds_token_ok = token and token == getattr(settings, "ONLYOFFICE_CALLBACK_TOKEN", "")
    if not ds_token_ok:
        if not request.user.is_authenticated:
            return HttpResponseForbidden()
        if not (work.is_archived or _can_access_work(request.user, work)):
            return HttpResponseForbidden()

    if not work.docx_file:
        if work.is_archived:
            return HttpResponseForbidden("DOCX missing")
        _generate_initial_docx(work, uploaded_by=request.user)
        if not work.docx_file:
            return HttpResponseForbidden("DOCX missing")

    resp = FileResponse(
        work.docx_file.open("rb"),
        as_attachment=False,
        filename=(work.docx_file.name or "work.docx").rsplit("/", 1)[-1],
    )
    return resp


@login_required
def onlyoffice_editor(request, work_id: int):
    work = get_object_or_404(Work, id=work_id)
    if not (work.is_archived or _can_access_work(request.user, work)):
        return render(request, "works/forbidden.html", status=403)

    if not work.docx_file:
        if work.is_archived:
            return render(
                request,
                "works/work_archive_no_docx.html",
                {"work": work},
                status=404,
            )
        _generate_initial_docx(work, uploaded_by=request.user)
        if not work.docx_file:
            return render(request, "works/forbidden.html", status=403)

    file_url, callback_url = _onlyoffice_file_and_callback_urls(request, work)

    can_edit = False
    can_comment = False
    mode = "edit"
    if work.is_archived:
        mode = "view"
    elif hasattr(request.user, "student") and work.author_id == request.user.student.id:
        can_edit = True
        can_comment = True
    elif hasattr(request.user, "teacher") and work.scientific_director_id == request.user.teacher.id:
        can_edit = True
        can_comment = True
    elif getattr(request.user, "user_type", None) == "department_admin":
        can_edit = True
        can_comment = True

    if work.is_archived:
        if request.GET.get("from") == "archive":
            onlyoffice_back_url = reverse("home")
        else:
            onlyoffice_back_url = reverse("work_detail", kwargs={"work_id": work.id})
    elif hasattr(request.user, "student") and work.author_id == request.user.student.id:
        onlyoffice_back_url = reverse("write_work")
    elif getattr(request.user, "user_type", None) in ("teacher", "department_admin") and _can_access_work(
        request.user, work
    ):
        onlyoffice_back_url = reverse("teacher_review_work", kwargs={"work_id": work.id})
    else:
        onlyoffice_back_url = reverse("home")

    doc_key_seed = f"work:{work.id}|t:{getattr(work, 'docx_updated_at', '')}"
    doc_key = hashlib.sha256(doc_key_seed.encode("utf-8")).hexdigest()[:32]

    editor_config = {
        "mode": mode,
        "lang": "ru",
        "customization": {
            "forcesave": bool(can_edit),
        },
        "user": {
            "id": str(request.user.id),
            "name": (request.user.get_full_name() or request.user.username),
        },
    }
    if can_edit:
        editor_config["callbackUrl"] = callback_url

    config = {
        "document": {
            "fileType": "docx",
            "key": doc_key,
            "title": f"{work.name}.docx",
            "url": file_url,
            "permissions": {
                "edit": can_edit,
                "comment": can_comment,
                "download": True,
                "print": True,
                "review": can_edit,
                "forcesave": bool(can_edit),
            },
        },
        "documentType": "word",
        "editorConfig": editor_config,
    }

    secret = getattr(settings, "ONLYOFFICE_JWT_SECRET", "")
    if secret:
        config["token"] = _jwt_hs256(config, secret)

    return render(
        request,
        "works/onlyoffice_editor.html",
        {
            "work": work,
            "onlyoffice_ds_url": settings.ONLYOFFICE_DS_URL,
            "onlyoffice_public_ds_url": settings.ONLYOFFICE_PUBLIC_DS_URL,
            "config": config,
            "onlyoffice_back_url": onlyoffice_back_url,
        },
    )


@csrf_exempt
def onlyoffice_callback(request, work_id: int):
    if request.GET.get("token") != getattr(settings, "ONLYOFFICE_CALLBACK_TOKEN", ""):
        return JsonResponse({"error": 1})

    work = get_object_or_404(Work, id=work_id)
    try:
        payload = json.loads(request.body.decode("utf-8"))
    except Exception:
        payload = {}

    status = payload.get("status")
    url = payload.get("url")

    if status in (2, 6) and url:
        if work.is_archived:
            return JsonResponse({"error": 0})
        try:
            ds_base = urlparse(getattr(settings, "ONLYOFFICE_DS_URL", "http://localhost:8082"))
            u = urlparse(url)
            if u.hostname and u.hostname not in {"127.0.0.1", "localhost"}:
                url = urlunparse((ds_base.scheme, ds_base.netloc, u.path, u.params, u.query, u.fragment))

            with urllib.request.urlopen(url) as r:                                         
                data = r.read()
            if work.docx_file:
                try:
                    work.docx_file.delete(save=False)
                except Exception:
                    pass
            work.docx_updated_at = timezone.now()
            work.docx_file.save(f"work_{work.id}_onlyoffice.docx", ContentFile(data), save=True)
            work.student_docx_reminder_sent_at = None
            work.teacher_stale_docx_notified_at = None
            work.save(
                update_fields=[
                    "student_docx_reminder_sent_at",
                    "teacher_stale_docx_notified_at",
                    "docx_updated_at",
                    "docx_file",
                ]
            )
        except Exception:
            return JsonResponse({"error": 1})

    return JsonResponse({"error": 0})
